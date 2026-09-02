from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
st = doc.styles['Normal']
st.font.name = 'Calibri'; st.font.size = Pt(10.5)
for s in ('Heading 1','Heading 2','Heading 3'):
    doc.styles[s].font.color.rgb = RGBColor(0x1F,0x38,0x64)
    doc.styles[s].font.name = 'Calibri'
sec = doc.sections[0]
sec.top_margin = sec.bottom_margin = Cm(2)
sec.left_margin = sec.right_margin = Cm(2.2)

def H(t, lvl=1): doc.add_heading(t, level=lvl)
def P(t='', bold=False, italic=False, size=None):
    p = doc.add_paragraph()
    r = p.add_run(t); r.bold = bold; r.italic = italic
    if size: r.font.size = Pt(size)
    return p
def B(t):
    doc.add_paragraph(t, style='List Bullet')
def N(t):
    doc.add_paragraph(t, style='List Number')
def T(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i,h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ''
        r = c.paragraphs[0].add_run(h); r.bold = True; r.font.size = Pt(9.5)
    for row in rows:
        cells = t.add_row().cells
        for i,v in enumerate(row):
            v = '' if v is None else str(v)
            lines = v.replace('\r','').split('\n')
            cells[i].text = ''
            para = cells[i].paragraphs[0]
            for j, ln in enumerate(lines):
                if j: para = cells[i].add_paragraph()
                rr_ = para.add_run(ln); rr_.font.size = Pt(9.5)
    if widths:
        for r_ in t.rows:
            for i,w in enumerate(widths):
                r_.cells[i].width = Cm(w)
    doc.add_paragraph()
    return t
def NOTE(t):
    p = doc.add_paragraph()
    r = p.add_run('Note   '); r.bold = True; r.font.size = Pt(9.5)
    r2 = p.add_run(t); r2.font.size = Pt(9.5); r2.italic = True

# ------------------------------------------------------------------ cover
h = doc.add_heading('Business Partner Mass Upload Programs', level=0)
P('User Acceptance Test Guide', bold=True, size=13)
P('Supplier master and Customer master mass upload for SAP S/4HANA', italic=True)
doc.add_paragraph()
T(['Item','Detail'],
  [['Programs','ZMMS_BP_MASS_UPLOAD  (supplier / vendor master)\nZSDS_CUST_MASS_UPLOAD  (customer master)'],
   ['System','CRS, client 500'],
   ['Release','SAP S/4HANA, S4CORE 109 / SAP_BASIS 816'],
   ['Source templates','Vendor LSMW with Template.xlsx\ncustomer master LSMW -  with format.xlsx'],
   ['Document purpose','Instructions for the business team to test both programs before go-live'],
   ['Prepared for','Testing team'],
  ], widths=[4.0,12.0])

doc.add_page_break()

# ------------------------------------------------------------------ 1
H('1.  Why these programs exist', 1)
P('Your existing upload templates were built as LSMW recordings of transactions XK01, XK02, XK05 and XD01. '
  'Under the SAP S/4HANA Business Partner approach those transactions are either redirected to transaction BP '
  'or removed altogether, so the recordings can no longer run.')
P('These two programs keep your existing spreadsheet layouts exactly as they are, and replace only the engine '
  'underneath. Nothing you already do in Excel has to change.')
P('Both programs create and change data through the standard SAP Business Partner interface, so every check, '
  'validation and number range that applies when a user works in transaction BP also applies here.')

# ------------------------------------------------------------------ 2
H('2.  Before you start testing', 1)
T(['You need','Why'],
  [['A test client with representative master data','So the checks against account groups, company codes and sales areas behave as they will in production'],
   ['Your normal BP creation authorisations','The programs run under your own user and enforce the same authorisation checks as transaction BP'],
   ['The two Excel workbooks','Use the same files you use today; no re-formatting is needed'],
   ['A few known-good customers / suppliers','For the change scenarios, so you can confirm existing data is not disturbed'],
  ], widths=[6.0,10.0])

H('2.1  The file layout - this applies to every tab of both workbooks', 2)
P('The programs find the data by its headings, not by its position.', bold=True)
P('Each scenario knows the headings that belong above its columns. On opening the file both programs look '
  'through every tab, and through the first ten lines of each tab, for the line carrying most of those '
  'headings. That line is the heading row; everything below it is data.')
P('What this means in practice:')
B('The tab may be called anything. You can send the whole workbook, or one tab saved on its own as '
  '"Sheet1" - both load the same way.')
B('The heading row does not have to be row 1. A title line above it is no longer a problem.')
B('Columns may be in a different order from the template, and columns you do not need may be deleted. '
  'Each column is read from wherever its heading actually is.')
B('A column heading may be either the wording used in the template ("Company Code") or the technical '
  'field name ("BUKRS"). Both are recognised.')
P('What still has to be right:', bold=True)
B('Please do not rename or translate the headings. A heading the program does not recognise makes that '
  'one column fall back to its position in the template, which is where a wrong value can creep in.')
B('If the same heading appears twice on a tab, neither copy can identify a column, so both fall back to '
  'their position. Please make repeated headings unique.')
B('Rows between the heading row and the first record - field type, field length, mandatory/optional, '
  'guideline and sample rows - must be deleted. They are read as data. Please delete them rather than '
  'hide them: a hidden row is still read.')
P('The companion Upload File Format Specification lists, tab by tab, every column and every row that has '
  'to go.')
NOTE('One line at the top of the result list tells you how many columns were found somewhere other than '
     'where the template has them. Nothing is wrong with that - it is there so you can see that the '
     'program noticed.')

# ------------------------------------------------------------------ 3
H('3.  The most important instruction: always test run first', 1)
P('Both programs have a Test run checkbox, and it is ticked by default. Leave it ticked for your first pass.')
P('In a test run the program does everything it would normally do, including every validation, but nothing '
  'is saved. You get the full result list showing exactly what would have happened, row by row.')
P('Only untick Test run once a file comes back with no errors.', bold=True)

H('3.1  Heading rows to skip', 2)
P('Both programs have a field called Heading rows to skip, set to 1. It is now only a fallback: it is used '
  'when no heading row can be recognised at all - for example a file sent with the heading row already '
  'removed. When the headings are there, the program finds them itself and this field is ignored.')
P('If a run reports that no tab carries the columns of the scenario, that is the case to look at: either '
  'the headings have been renamed, or the file is the wrong one for the radio button selected.')

doc.add_page_break()

# ------------------------------------------------------------------ 4
H('4.  Program A — Supplier master upload', 1)
P('Program name: ZMMS_BP_MASS_UPLOAD', bold=True)
H('4.1  Selection screen', 2)
T(['Field','What to enter'],
  [['Scenario','Pick one of the nine radio buttons. This decides which tab of the workbook is read.'],
   ['Upload workbook','The path to Vendor LSMW with Template.xlsx'],
   ['File is on the PC / application server','Where the file sits. Use PC for normal testing.'],
   ['Test run (nothing is posted)','Leave ticked for the first pass'],
   ['Stop at the first faulty row','Tick this if you want the run to halt as soon as something fails, instead of working through the whole file'],
   ['Heading rows to skip','Leave at 1. See section 3.1.'],
  ], widths=[5.5,10.5])

H('4.2  The nine scenarios', 2)
T(['Radio button','Workbook tab','What it does'],
  [['Vendor / BP creation - all company codes','Vendor creation for All CC','Creates the supplier and its Business Partner, with company code and purchasing data'],
   ['Withholding tax / TDS','TDS upload','Maintains withholding tax types and codes per company code'],
   ['TAN exemption details','TAN details','Maintains the India TAN exemption records'],
   ['Bank key creation','BANK Key creation','Creates or changes bank master records'],
   ['Vendor bank details','Bank details update','Maintains the supplier’s own bank accounts'],
   ['Vendor extension','Vendor extension','Extends an existing supplier to a further company code or purchasing organisation'],
   ['CIN details','CIN details','Maintains the India tax and excise fields'],
   ['Partner functions','Patner function','Maintains purchasing partner functions'],
   ['Block / unblock','Block_Unblocked','Sets or clears posting and purchasing blocks'],
  ], widths=[4.8,4.2,7.0])
NOTE('The tab names in this table are the ones in your workbook - the spelling "Patner function" included. They are shown so you can see which tab each radio button was built from; the program no longer depends on them, so renaming a tab does no harm.')

doc.add_page_break()

# ------------------------------------------------------------------ 5
H('5.  Program B — Customer master upload', 1)
P('Program name: ZSDS_CUST_MASS_UPLOAD', bold=True)
H('5.1  Selection screen', 2)
T(['Field','What to enter'],
  [['Scenario','Pick one of the seven radio buttons'],
   ['Upload workbook','The path to customer master LSMW -  with format.xlsx'],
   ['File is on the PC / application server','Where the file sits'],
   ['Test run (nothing is posted)','Leave ticked for the first pass'],
   ['Stop at the first faulty row','Halt on the first failure instead of processing the whole file'],
   ['BP grouping (blank = derived)','Leave this blank. SAP then derives the BP grouping from the account group, which is the normal behaviour. Only fill it if you need to force a particular grouping.'],
   ['Heading rows to skip','Leave at 1. See section 3.1.'],
  ], widths=[5.5,10.5])

H('5.2  The seven scenarios', 2)
T(['Radio button','Workbook tab','Columns','What it does'],
  [['Domestic customer - India','domestic customer IND','136','Creates Indian domestic customers with company code, sales area, tax and licence data'],
   ['Export customer','Export customer','67','Creates export customers'],
   ['Morocco customer','Morocco customer','116','Creates Morocco customers'],
   ['SAGA customer','SAGA customer','122','Creates SAGA customers'],
   ['Credit limit (FSCM)','credit Limit','18','Sets credit limits, risk class and credit blocks'],
   ['Domestic customer - US','domestic customer US','75','Creates US domestic customers'],
   ['Ship-to party - US','ship to party US','75','Creates US ship-to parties'],
  ], widths=[4.2,4.0,1.8,6.0])

doc.add_page_break()

# ------------------------------------------------------------------ 6
H('6.  How to run a test', 1)
N('Start the program in transaction SE38 or through the transaction code once it is assigned.')
N('Choose the scenario that matches the tab you want to load.')
N('Select the workbook.')
N('Leave Test run ticked.')
N('Execute.')
N('The program locates the heading row and processes every row below it.')
N('A result list appears. Work through any errors, correct the spreadsheet, and run again.')
N('When the test run is clean, untick Test run and execute again to post for real.')

H('6.1  Reading the result list', 2)
P('Every row of your spreadsheet produces at least one line in the result list.')
T(['Column','Meaning'],
  [['Status','Green = fine, yellow = warning, red = error'],
   ['Excel row','The row number in your spreadsheet, so you can go straight to it'],
   ['Customer / Vendor','The account the message relates to'],
   ['Structure and Field','Which field caused the problem. This is the most useful column when something fails — it points at the spreadsheet column rather than giving a general message.'],
   ['Message','The SAP message text'],
  ], widths=[4.0,12.0])
P('A summary line at the bottom of the screen tells you how many rows were read, how many were processed '
  'and how many had errors.')

# ------------------------------------------------------------------ 7
H('7.  What to check after a live posting', 1)
T(['What you loaded','Where to check it'],
  [['Supplier or customer, general data','Transaction BP — General Data'],
   ['Company code data','Transaction BP — role FLCU00 (customer) or FLVN00 (supplier)'],
   ['Sales area data','Transaction BP — role FLCU01, Sales and Distribution'],
   ['Purchasing data','Transaction BP — role FLVN01'],
   ['Address, telephone, fax, e-mail','Transaction BP — Address tab. Mobile numbers appear flagged as mobile.'],
   ['Tax classifications','Transaction BP — Sales and Distribution, Billing, Tax classification'],
   ['Licences, bank guarantee, routing','Transaction BP — role FLCU01, tab "Licenses and Their Validity"'],
   ['Credit limits and risk class','Transaction UKM_BP, or BP in role UKM000'],
   ['Bank master records','Transaction FI03'],
  ], widths=[6.0,10.0])

doc.add_page_break()

# ------------------------------------------------------------------ 8
H('8.  Suggested test cases', 1)
P('We suggest working through these in order. Each one is quick and isolates a different part of the program.')
T(['#','Test','How','What should happen'],
  [['1','Empty file','Run any scenario against a tab with only the heading row','A message saying the tab holds no data below its heading. Nothing is posted.'],
   ['2','Wrong file','Point the program at the customer workbook while a supplier scenario is selected','A clear message that no tab in that workbook carries the columns of the selected scenario'],
   ['3','One good row, test run','One valid new customer or supplier','Green line, "Test run OK". Check in BP that nothing was actually created.'],
   ['4','One good row, live','Same row with Test run unticked','Green line, record created. Verify it in BP.'],
   ['5','Deliberate error','Put an account group that does not exist in the account group column','Red line naming the field and the value. Nothing posted for that row.'],
   ['6','Bad date','Type 31.02.2026 in a date column','Red line saying the value is not a valid date'],
   ['7','Stop at first error','Two bad rows, tick "Stop at the first faulty row"','The run halts after the first failure and says so'],
   ['8','Change an existing record','Fill only the key columns and one field you want to change','Only that field changes. Everything else stays as it was — see section 9.'],
   ['9','Mixed file','Ten rows, three of them faulty','Seven green, three red. The seven are posted; the three are not.'],
   ['10','Licence data','A customer row with drug licence and bank guarantee columns filled','Values appear on the "Licenses and Their Validity" tab in BP'],
   ['11','Credit limit','A row on the credit Limit tab','Limit visible in UKM_BP for the matching credit segment'],
  ], widths=[0.9,3.4,5.2,6.5])

doc.add_page_break()

# ------------------------------------------------------------------ 9
H('9.  Behaviour you should expect', 1)
H('9.1  A blank cell does not erase data', 2)
P('If you leave a cell empty, the program leaves whatever is already in SAP untouched. This means you can '
  'load a file that only fills the columns you actually want to change.')
P('If you genuinely want to clear a field, type the word #BLANK# in the cell.', bold=True)

H('9.2  One licence record per customer', 2)
P('The licence and bank guarantee data is stored against the customer as a whole. The Plant field on that '
  'screen is informational and is not part of the key, which means a customer can hold exactly one set of '
  'licence details. If a spreadsheet carries the licence block more than once, only the first set can be loaded.')

H('9.3  Credit limit currency is checked, not written', 2)
P('In SAP S/4HANA a credit limit is held in the currency of its credit segment, which is set in configuration. '
  'The currency column on the credit Limit tab is therefore validated against the segment and the row is '
  'rejected if the two disagree — but the currency itself is not written.')
P('In your system the credit segment is the same code as the credit control area. For example credit control '
  'area 1000 uses segment 1000 in INR, 7450 uses segment 7450 in EUR, and 6600 uses segment 6600 in MAD.')

H('9.4  What the credit Limit tab writes, and where', 2)
P('The tab carries three fields that are customer master data rather than credit data: payment terms, the '
  'interest indicator and customer group 3. They are written through the same Business Partner interface as '
  'every other scenario - nothing is written to a table directly.')
P('Payment terms and the interest indicator belong to a company code, and customer group 3 to a sales area, '
  'but the tab carries neither. The program therefore uses the ones the customer already has:')
T(['Situation','What happens'],
  [['The customer is in one company code','Payment terms and interest indicator are written there'],
   ['The customer is in several company codes','The company codes belonging to the credit control area in '
    'the row are used to narrow it down. If exactly one remains, that is used'],
   ['It still cannot be narrowed to one','A red line naming the company codes, and the two fields are not '
    'written. The credit limit itself is unaffected. Add a company code column if you need this'],
   ['The customer has one sales area','Customer group 3 is written there'],
   ['The customer has several sales areas','A red line, and customer group 3 is not written']],
  widths=[6.0,10.0])

H('9.5  Two field limits in the existing licence table', 2)
T(['Field','Limit','Effect'],
  [['Bank Guarantee Amount','Whole numbers only, maximum 10 digits','A guarantee of 2,500,000.50 is stored as 2,500,000. If decimals are needed the field itself has to be changed.'],
   ['Distance in kms','Maximum 32,767','A larger value cannot be stored and the row is rejected'],
  ], widths=[4.0,4.0,8.0])

H('9.6  Tax classifications depend on the country', 2)
P('The tax classification columns are filled according to the tax categories configured for the country in question:')
T(['Country','Tax categories configured'],
  [['India','JTX1, JTX2, JTX3, JTX4, JOCG, JTC1'],
   ['United States','UTXJ, UTX2, UTX3, MWST'],
   ['Spain','MWST only'],
   ['Morocco','ZMVT only'],
  ], widths=[4.0,12.0])
P('This is why some tax classification columns on the Morocco and SAGA tabs cannot be loaded — see section 10.')

doc.add_page_break()

# ------------------------------------------------------------------ 10
H('10.  Corrections we applied to the templates', 1)
P('While mapping the spreadsheets we found several columns that could not be loaded as labelled. '
  'We corrected them rather than let them fail silently. Please review and confirm.')
T(['Tab','Column','What we found','What the program does'],
  [['domestic customer IND','BY','The heading reads JOIG, which is Integrated GST, but the description says Central GST. Your configuration has JOCG, not JOIG.','Loads the column as JOCG'],
   ['Morocco customer','61–64','TAXKD_02 to TAXKD_05, but Morocco has only one tax category (ZMVT)','The four columns are not loaded'],
   ['SAGA customer','72–75','TAXKD_02 to TAXKD_05, but Spain has only one tax category (MWST)','The four columns are not loaded'],
   ['Morocco customer','112–116','Repeat description text for columns 107–111 rather than real fields','Not loaded'],
   ['SAGA customer','36','No heading at all. Both your own LSMW recording and the India tab place the vendor number here.','Loaded as the vendor number'],
   ['SAGA customer','39–41','The three Spanish DIR3 codes required for FACe e-invoicing, but two of them pointed at the same field','Accounting Office to Tax Number 3, Managing Office to Tax Number 4, Processing Unit to Tax Number 5'],
   ['SAGA customer','44','Repeats Tax Number 3, but is annotated with a reconciliation account','Loaded as the reconciliation account'],
   ['credit Limit','14','Credit representative group. This is not configured in your system — the field is empty for every credit control area.','Not loaded'],
   ['credit Limit','R','Holds two values in one column: the interest indicator and the interest cycle','Loaded as both - the first value becomes the interest indicator, a number after it the calculation cycle in months. Splitting the column is still cleaner but no longer necessary.'],
  ], widths=[3.0,1.6,6.4,5.0])

# ------------------------------------------------------------------ 11
H('11.  Points we still need from you', 1)
T(['#','Point','What we need'],
  [['1','Aadhaar number','SAP has no standard field for Aadhaar. We propose the identification type X90003, matching the X90001 and X90002 types you already use. Please confirm it has been created.'],
   ['2','SAGA tab, Spanish DIR3 codes','Please confirm the allocation in section 10. Also note that columns CG and CH are annotated "VAT number" and "Accounting Office" but are in fact the 20B and 21B drug licence fields — loading Spanish tax data there would overwrite licence values that your invoice and credit note programs read.'],
   ['3','SAGA customer and domestic customer IND tabs','No longer blocking - the program finds the heading row on row 2 by itself. The descriptive rows between the headings and the first record still have to go.'],
   ['4','Sample records','Two existing customers that already carry licence data and a bank guarantee, one Indian and one Morocco or SAGA, so we can verify the mapping against real values.'],
  ], widths=[0.9,3.6,11.5])

# ------------------------------------------------------------------ 12
H('12.  If something goes wrong', 1)
P('Please send us the following, and we can usually identify the cause without needing access to your session:')
B('The spreadsheet you used, or at least the failing rows')
B('Which scenario radio button was selected')
B('Whether Test run was ticked')
B('A screenshot of the result list, including the Structure and Field columns')
B('The customer or supplier number, if one was created')
doc.add_paragraph()
NOTE('If a row fails, nothing from that row is saved. The remaining rows are unaffected, so it is always safe to correct the file and run it again.')

doc.save('docs/BP_Mass_Upload_Test_Guide.docx')
print('written')
