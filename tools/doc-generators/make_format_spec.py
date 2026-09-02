import zipfile,json,re
from xml.etree import ElementTree as ET
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT

NS='{http://schemas.openxmlformats.org/spreadsheetml/2006/main}'
R='{http://schemas.openxmlformats.org/officeDocument/2006/relationships}'
def letters(n):
    s=''
    while n: n,r=divmod(n-1,26); s=chr(65+r)+s
    return s
def rowsof(path,upto=12):
    z=zipfile.ZipFile(path)
    wb=ET.fromstring(z.read('xl/workbook.xml'))
    rels={r.get('Id'):r.get('Target') for r in ET.fromstring(z.read('xl/_rels/workbook.xml.rels'))}
    ss=[''.join(t.text or '' for t in si.iter(NS+'t')) for si in ET.fromstring(z.read('xl/sharedStrings.xml'))]
    def cn(ref):
        s=''.join(c for c in ref if c.isalpha()); n=0
        for c in s: n=n*26+ord(c)-64
        return n
    out={}
    for sh in wb.iter(NS+'sheet'):
        nm=sh.get('name'); t=rels[sh.get(R+'id')]; t=t if t.startswith('xl/') else 'xl/'+t
        w=ET.fromstring(z.read(t)); rr={}
        for row in w.iter(NS+'row'):
            r=int(row.get('r'))
            if r>upto: break
            d={}
            for c in row.iter(NS+'c'):
                v=c.find(NS+'v'); x='' if v is None else v.text
                if c.get('t')=='s': x=ss[int(x)]
                x=(x or '').strip()
                if x: d[cn(c.get('r'))]=x
            rr[r]=d
        out[nm]=rr
    return out
VEN=rowsof('/home/user/ONGC-CST-Purchase-Date-Sharing-/Vendor LSMW with Template.xlsx')
D=json.load(open('/tmp/claude-0/-home-user-ONGC-CST-Purchase-Date-Sharing-/0a870517-dcc5-5a06-9069-d731d41a85f0/scratchpad/ddic2.json'))
DLEN=json.load(open('/tmp/claude-0/-home-user-ONGC-CST-Purchase-Date-Sharing-/0a870517-dcc5-5a06-9069-d731d41a85f0/scratchpad/ddic.json'))
prog=open('/home/user/ONGC-CST-Purchase-Date-Sharing-/src/zsds_cust_mass_upload.prog.abap').read()
MAP=re.findall(r"scen = '(\w+)' col = (\d+)\s+node = '(\w)' fld = '([^']*)' cnv = '(\w*)' \)(?:\s+\" (.*))?", prog)

doc=Document(); s=doc.styles['Normal']; s.font.name='Calibri'; s.font.size=Pt(9.5)
for h in ('Heading 1','Heading 2','Heading 3'):
    doc.styles[h].font.color.rgb=RGBColor(0x1F,0x38,0x64); doc.styles[h].font.name='Calibri'
sec=doc.sections[0]; sec.top_margin=sec.bottom_margin=Cm(1.6); sec.left_margin=sec.right_margin=Cm(1.6)
def H(t,l=1): doc.add_heading(t,level=l)
def P(t='',bold=False,italic=False,size=None):
    p=doc.add_paragraph(); r=p.add_run(t); r.bold=bold; r.italic=italic
    if size: r.font.size=Pt(size)
def B(t): doc.add_paragraph(t,style='List Bullet')
def T(hdr,rows,widths=None,fs=8.5):
    t=doc.add_table(rows=1,cols=len(hdr)); t.style='Light Grid Accent 1'
    t.alignment=WD_TABLE_ALIGNMENT.LEFT
    for i,h in enumerate(hdr):
        c=t.rows[0].cells[i]; c.text=''
        r=c.paragraphs[0].add_run(h); r.bold=True; r.font.size=Pt(fs)
    for row in rows:
        cs=t.add_row().cells
        for i,v in enumerate(row):
            v = '' if v is None else str(v)
            lines = v.replace('\r','').split('\n')
            cs[i].text = ''
            para = cs[i].paragraphs[0]
            for j, ln in enumerate(lines):
                if j: para = cs[i].add_paragraph()
                rr_ = para.add_run(ln); rr_.font.size = Pt(fs)
    if widths:
        for rr_ in t.rows:
            for i,w in enumerate(widths): rr_.cells[i].width=Cm(w)
    doc.add_paragraph(); return t
def NOTE(t):
    p=doc.add_paragraph(); r=p.add_run('Note   '); r.bold=True; r.font.size=Pt(8.5)
    r2=p.add_run(t); r2.italic=True; r2.font.size=Pt(8.5)
json.dump({'ok':1},open('/dev/null','w')) if False else None

# ------------------------------------------------------------------ layouts
# tab -> (tech row, description row, type row, length row, M/O row, first data row)
VLAY={
 'Vendor creation for All CC': (1,2,None,None,3,4),
 'TDS upload':                 (1,2,None,None,None,3),
 'TAN details':                (1,None,None,None,None,2),
 'BANK Key creation':          (1,5,2,3,4,7),
 'Bank details update':        (1,5,2,3,4,7),
 'Vendor extension':           (1,5,2,3,4,7),
 'CIN details':                (1,None,None,None,None,2),
 'Patner function':            (6,9,7,8,None,10),
 'Block_Unblocked':            (4,7,5,6,None,9),
}
VSCEN=[('Vendor / BP creation - all company codes','Vendor creation for All CC'),
       ('Withholding tax / TDS','TDS upload'),
       ('TAN exemption details','TAN details'),
       ('Bank key creation','BANK Key creation'),
       ('Vendor bank details','Bank details update'),
       ('Vendor extension','Vendor extension'),
       ('CIN details','CIN details'),
       ('Partner functions','Patner function'),
       ('Block / unblock','Block_Unblocked')]
CSCEN=[('R1','Domestic customer - India','domestic customer IND',2,3),
       ('R2','Export customer','Export customer',1,4),
       ('R3','Morocco customer','Morocco customer ',1,4),
       ('R4','SAGA customer','SAGA customer',2,4),
       ('R5','Credit limit (FSCM)','credit Limit',1,4),
       ('R6','Domestic customer - US','domestic customer US',1,2),
       ('R7','Ship-to party - US','ship to party US',1,2)]
NODE={'K':'Key field','A':'Address','M':'Communication','C':'Customer general data',
      'B':'Company code data','S':'Sales area data','T':'Tax classification',
      'Z':'Licences and Their Validity','I':'BP identification','U':'Credit management'}
SRC={'C':'KNA1','B':'KNB1','S':'KNVV','Z':'ZSD_LICENSE_CHK','A':'ADRC'}
def fmt_of(node,fld):
    tab=SRC.get(node)
    if not tab: return ''
    for c in D.get(tab,[]):
        if c['f']==fld:
            dt=c['dt'] or ''
            ln=(DLEN.get(tab,{}).get(fld,{}) or {}).get('len','')
            nice={'CHAR':'Text','DATS':'Date','NUMC':'Digits','CUKY':'Currency',
                  'INT2':'Whole number','CLNT':'Client','DEC':'Number','CURR':'Amount'}.get(dt,dt)
            return f'{nice} {ln}'.strip()
    return ''

# ------------------------------------------------------------------ cover
doc.add_heading('Upload File Format Specification',level=0)
P('Column-by-column layout for every scenario of both mass upload programs',bold=True,size=12)
P('Companion to the User Acceptance Test Guide',italic=True)
doc.add_paragraph()
T(['Programs','Scenarios','Columns documented'],
  [['ZMMS_BP_MASS_UPLOAD  (supplier master)','9','238'],
   ['ZSDS_CUST_MASS_UPLOAD  (customer master)','7','607']],widths=[8.0,3.0,5.0],fs=9.5)

H('1.  How to read this document',1)
P('There is one section per radio button. Each section names the workbook tab it was built from, states how '
  'that tab is laid out today, and lists every column.')
P('The programs find the data by its headings, not by its position.',bold=True)
P('Each scenario knows the headings that belong above its columns. On opening the file both programs look '
  'through every tab of the workbook, and through the first ten lines of each tab, for the line carrying '
  'most of those headings. That line is the heading row; everything below it is data. So:')
B('The tab may be called anything, and the workbook may hold one tab or all of them.')
B('The heading row does not have to be row 1.')
B('Columns may be in a different order from the template, and columns you do not need may be deleted - '
  'each column is read from wherever its heading actually is.')
B('A heading may be the wording used in the template ("Company Code") or the technical field name '
  '("BUKRS"). Both are recognised.')
B('A heading used more than once on a tab is matched by its turn - the second one in the file is the '
  'second one in the template - as long as the file repeats it at least as often as the template does.')
B('Headings running across two lines (a technical name on one, a description on the next, as on the '
  'credit tab) are both used.')
P('What still has to be right:',bold=True)
B('Headings must not be renamed or translated. A column whose heading is not recognised is read from the '
  'position it has in the template, which is only correct if that part of the file matches the template.')
B('Such a column is left empty rather than loaded when its position turns out to hold a different field. '
  'The result list names those columns.')
B('Every row between the heading row and the first record must be deleted: field type, field length, '
  'mandatory/optional, guideline and sample rows are all read as data. Please delete them rather than '
  'hide them - a hidden row is still read.')
P('Each section below states exactly which rows to delete from that tab.')
P('The column letters in the tables below are the letters the template uses. They are there so you can '
  'find a column, not as a requirement: if a letter does not match your file, the program still reads the '
  'column by its heading. One line at the top of the result list says how many columns were found '
  'somewhere else.')

H('2.  Heading rows to skip',1)
P('Both programs have a field on the selection screen called Heading rows to skip, set to 1. It is now '
  'only a fallback, used when no heading row can be recognised at all - a file sent with the heading row '
  'already removed, for instance. When the headings are present the program finds them itself and the '
  'field is ignored.')
P('If a run reports that no tab in the workbook carries the columns of the scenario, either the headings '
  'have been renamed or the file does not match the radio button selected.')

H('3.  Mandatory and optional',1)
P('Where the workbook itself marks a column Mandatory or Optional, that marking is reproduced here. Where it '
  'does not, the column is shown blank — SAP still applies its own required-field rules for the account group, '
  'so a field can be rejected at posting even if it is not marked mandatory in the spreadsheet.')
P('A blank cell never erases data.',bold=True)
P('If you leave a cell empty the program leaves whatever is already in SAP untouched. To clear a field on '
  'purpose, type #BLANK# in the cell.')
doc.add_page_break()

# ------------------------------------------------------------------ vendor
H('4.  Supplier master - ZMMS_BP_MASS_UPLOAD',1)
n=0
for radio,tab in VSCEN:
    n+=1
    H(f'4.{n}  {radio}',2)
    tech,desc,typ,lng,mo,first = VLAY[tab]
    rr=VEN[tab]
    P(f'Workbook tab this was built from:  {tab}   (the tab may be renamed - the program does not read the name)',bold=True)
    # Only the rows BETWEEN the headings and the first record have to go -
    # anything above the heading row is stepped over by the program itself.
    named={}
    if desc: named[desc]='field description'
    if typ:  named[typ]='field type'
    if lng:  named[lng]='field length'
    if mo:   named[mo]='mandatory/optional'
    gone=list(range(tech+1,first))
    what=[named[r] for r in gone if r in named]
    if [r for r in gone if r not in named]: what.append('sample and guideline rows')
    ok = not gone
    todo = ('Nothing - this tab is already correct' if ok
            else (f'Delete row{"s" if len(gone)>1 else ""} {gone[0]}'
                  + (f' to {gone[-1]}' if len(gone)>1 else '')
                  + (' (' + ', '.join(what) + ')' if what else '')
                  + f', so the first record sits directly under the heading row on row {tech}'))
    T(['Required','Detail'],
      [['Heading row','anywhere in the first ten rows - the program finds it'],
       ['First data row','the row directly under the headings'],
       ['Headings are on','row %d today' % tech],
       ['First record is on','row %d today' % first],
       ['To do',todo]],widths=[4.0,12.0],fs=9)
    mx=max((max(d) if d else 0) for d in rr.values())
    rows=[]
    for i in range(1,mx+1):
        t_=rr.get(tech,{}).get(i,'')
        if not t_: continue
        rows.append([letters(i), t_,
                     rr.get(desc,{}).get(i,'') if desc else '',
                     rr.get(typ,{}).get(i,'') if typ else '',
                     rr.get(lng,{}).get(i,'') if lng else '',
                     rr.get(mo,{}).get(i,'') if mo else ''])
    T(['Col','Technical name','Description','Type','Len','M/O'],rows,
      widths=[1.0,3.2,7.4,1.6,1.0,1.0])
    doc.add_page_break()

# ------------------------------------------------------------------ customer
H('5.  Customer master - ZSDS_CUST_MASS_UPLOAD',1)
CUSROWS=rowsof('/home/user/ONGC-CST-Purchase-Date-Sharing-/customer master LSMW -  with format.xlsx')
n=0
for scen,radio,tab,namerow,first in CSCEN:
    n+=1
    H(f'5.{n}  {radio}',2)
    rr=CUSROWS[tab]
    P(f'Workbook tab this was built from:  {tab.strip()}   (the tab may be renamed - the program does not read the name)',bold=True)
    gone=list(range(namerow+1,first))
    ok = not gone
    todo=('Nothing - this tab is already correct' if ok
          else (f'Delete row{"s" if len(gone)>1 else ""} {gone[0]}'
                + (f' to {gone[-1]}' if len(gone)>1 else '')
                + f', so the first record sits directly under the heading row on row {namerow}'))
    T(['Required','Detail'],
      [['Heading row','anywhere in the first ten rows - the program finds it'],
       ['First data row','the row directly under the headings'],
       ['Headings are on',f'row {namerow} today'],
       ['First record is on',f'row {first} today'],
       ['To do',todo]],widths=[4.0,12.0],fs=9)
    ent={int(c):(nd,fl,cv,(cm or '').strip()) for sc,c,nd,fl,cv,cm in MAP if sc==scen}
    mx=max(list(ent)+[max(d) if d else 0 for d in rr.values()])
    rows=[]
    for i in range(1,mx+1):
        head=rr.get(namerow,{}).get(i,'')
        if i in ent:
            nd,fl,cv,cm=ent[i]
            goes=NODE.get(nd,'')
            det=fl if nd not in ('T',) else (f'tax category {fl}' if not fl.startswith('#') else f'tax category no. {fl[1:]} for the country')
            f_=fmt_of(nd,fl) or ('Date' if cv=='DT' else ('Digits' if cv=='NM' else ''))
            rows.append([letters(i), head or cm, goes, det, f_])
        else:
            if head: rows.append([letters(i), head, 'not loaded', '', ''])
    T(['Col','Heading in the workbook','Goes to','Field','Format'],rows,
      widths=[1.0,5.4,3.4,3.2,2.2])
    if scen=='R5':
        P('Credit limits are written through SAP Credit Management (FSCM): the limit is set on the credit '
          'segment of the credit control area in the row, the total limit on the main segment 0000, the risk '
          'class on the business partner and the block on the segment. Nothing is written to the old credit '
          'management tables, which S/4HANA no longer maintains.')
        P('Payment terms, interest indicator and customer group 3 on this tab are customer master data, not '
          'credit data. They are written through the Business Partner interface, into the company code and '
          'sales area the customer already has - see section 9.4 of the test guide. The interest column may '
          'hold the indicator and the calculation cycle together ("Z1 3"); both are loaded.')
    doc.add_page_break()

doc.save('/home/user/ONGC-CST-Purchase-Date-Sharing-/docs/BP_Mass_Upload_File_Format_Spec.docx')
print('saved')
